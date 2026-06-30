from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parents[1] / "deliverables" / "PROJECT REPORT - KIENCARE COMPLETED.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK = "1F4D78"
LIGHT = "F2F4F7"
MID = "D9E2F3"
TEXT = "1F2937"
MUTED = "5B6573"
GREEN = "E2F0D9"
AMBER = "FFF2CC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[idx]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def style_run(run, size=10.5, bold=False, color=TEXT, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_p(doc, text="", bold_prefix=None, italic=False, align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        style_run(p.add_run(bold_prefix), bold=True)
        style_run(p.add_run(text[len(bold_prefix):]), italic=italic)
    else:
        style_run(p.add_run(text), italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    style_run(p.add_run(text))
    return p


def new_numbering(doc):
    numbering = doc.part.numbering_part.element
    ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abs_id = max(ids, default=-1) + 1
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abs_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, attr, value in [("w:start", "w:val", "1"), ("w:numFmt", "w:val", "decimal"), ("w:lvlText", "w:val", "%1."), ("w:lvlJc", "w:val", "left")]:
        n = OxmlElement(tag)
        n.set(qn(attr), value)
        lvl.append(n)
    pPr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    pPr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    pPr.append(ind)
    lvl.append(pPr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    aid = OxmlElement("w:abstractNumId")
    aid.set(qn("w:val"), str(abs_id))
    num.append(aid)
    numbering.append(num)
    return num_id


def add_number(doc, text, num_id):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    style_run(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_table(doc, headers, rows, widths, fills=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        style_run(r, size=font_size, bold=True, color="FFFFFF")
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if fills and r_idx < len(fills) and fills[r_idx]:
                set_cell_shading(cells[i], fills[r_idx])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            style_run(p.add_run(str(value)), size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

# Standard business brief tokens.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(TEXT)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for level, size, color, before, after in [(1, 16, BLUE, 16, 8), (2, 13, BLUE, 12, 6), (3, 12, DARK, 8, 4)]:
    st = doc.styles[f"Heading {level}"]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

# Running header/footer.
hp = sec.header.paragraphs[0]
hp.text = "KIENCARE / LAPTOPHUB  |  TECHNICAL PROJECT REPORT"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
style_run(hp.runs[0], size=8.5, bold=True, color=MUTED)
add_page_number(sec.footer.paragraphs[0])

# Cover: editorial report pattern, restrained to the selected preset.
doc.add_paragraph().paragraph_format.space_after = Pt(78)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
style_run(p.add_run("TECHNICAL PROJECT REPORT"), size=12, bold=True, color=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
style_run(p.add_run("KIENCARE / LAPTOPHUB"), size=30, bold=True, color=DARK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(28)
style_run(p.add_run("Mobile E-commerce Application for Laptop Sales"), size=15, color=MUTED)
add_table(doc, ["Thông tin", "Nội dung"], [
    ["Môn học", "Flutter Application Development (PRM393)"],
    ["Loại dự án", "Ứng dụng thương mại điện tử trên thiết bị di động"],
    ["Nhóm", "4 thành viên"],
    ["Phiên bản kiểm chứng", "Source code tại workspace ngày 22/06/2026"],
    ["Nền tảng mục tiêu", "Android (Flutter); có cấu trúc đa nền tảng"],
], [2100, 7260], font_size=10)
add_p(doc, "Báo cáo được hoàn thiện từ yêu cầu/template cung cấp và được đối chiếu trực tiếp với mã nguồn, cấu hình, kiểm thử và kết quả build thực tế của dự án.", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
doc.add_page_break()

add_heading(doc, "MỤC LỤC", 1)
toc = [
    "1. Team Introduction", "2. Case Study", "3. Business Analysis / System Design",
    "4. Development Requirements", "5. Demo of Functions", "6. Conclusion and Discussion",
    "7. Contribution Table", "8. References", "Phụ lục A. Ma trận đối chiếu với source code"
]
for item in toc:
    add_p(doc, item, after=4)
add_p(doc, "Lưu ý: số trang được hiển thị tự động ở chân trang; mục lục này là danh mục nội dung tĩnh để bảo đảm hiển thị ổn định trên Word và LibreOffice.", italic=True)

add_heading(doc, "1. Team Introduction", 1)
add_p(doc, "Nhóm gồm bốn thành viên. Vai trò dưới đây được tổng hợp từ phạm vi file và lịch sử commit hiện có trong Git; tỷ lệ đóng góp được trình bày cân bằng theo trách nhiệm chung của đồ án.")
add_table(doc, ["Thành viên", "Vai trò chính", "Phạm vi đã triển khai", "Đóng góp"], [
    ["Bùi Mạnh Thịnh", "Backend/Data & Firebase", "Mô hình và dữ liệu sản phẩm, repository xác thực, cấu hình Firebase, nền tảng chat và tích hợp dữ liệu", "25%"],
    ["Lê Trung Kiên", "Architecture, UI/UX & Admin", "Kiến trúc ứng dụng, điều hướng, giao diện chính, bản đồ, chat và bộ màn hình quản trị", "25%"],
    ["Trần Đình Việt", "Checkout & Order Flow", "Giỏ hàng, địa chỉ, checkout, phương thức thanh toán, lịch sử/chi tiết đơn và quản trị đơn", "25%"],
    ["Trần Văn Phi", "Wishlist & Engagement", "SQLite wishlist, provider yêu thích, promotion, home, notification và các thành phần tương tác", "25%"],
], [1700, 1900, 4660, 1100], font_size=8.5)

add_heading(doc, "2. Case Study", 1)
add_p(doc, "Project title: KienCare / LaptopHub - Ứng dụng mua sắm laptop trên thiết bị di động.", bold_prefix="Project title:")
add_p(doc, "Domain: Online Sales System (E-commerce).", bold_prefix="Domain:")
add_p(doc, "KienCare/LaptopHub hỗ trợ khách hàng khám phá và mua laptop theo nhóm Gaming, Văn phòng, Doanh nhân và Laptop 2-in-1. Dữ liệu hiện có 28 sản phẩm từ tệp JSON, bao gồm thông số CPU/GPU/RAM/SSD/màn hình, giá, tồn kho, đánh giá và từ khóa tìm kiếm. Người dùng có thể xem sản phẩm, tìm kiếm/lọc/sắp xếp, lưu yêu thích, quản lý giỏ, nhập địa chỉ, đặt hàng, chọn COD/Momo/VNPay, xem lịch sử đơn, nhận thông báo, chat hỗ trợ và tìm cửa hàng trên Google Maps.")
add_p(doc, "Phạm vi quản trị gồm dashboard, CRUD sản phẩm, quản lý đơn, doanh thu và tài khoản. Tuy nhiên, phần lớn repository nghiệp vụ hiện là dữ liệu mô phỏng hoặc lưu cục bộ; riêng chat có thể dùng Cloud Firestore khi Firebase khởi tạo thành công. Vì vậy đây là prototype chức năng hoàn chỉnh cho demo học phần, chưa phải hệ thống thương mại điện tử production.")

add_heading(doc, "2.1. Problem Statement", 2)
add_bullet(doc, "Người mua laptop khó so sánh nhanh cấu hình, giá và nhóm nhu cầu trên giao diện di động.")
add_bullet(doc, "Một quy trình mua sắm cần duy trì giỏ hàng/đơn hàng cục bộ, hỗ trợ yêu thích và cung cấp kênh hỗ trợ trực tiếp.")
add_bullet(doc, "Người quản trị cần một khu vực tổng quan để theo dõi sản phẩm, đơn hàng, doanh thu và người dùng trong cùng ứng dụng demo.")

add_heading(doc, "2.2. Project Objectives", 2)
objective_num = new_numbering(doc)
for x in [
    "Xây dựng ứng dụng Flutter theo kiến trúc phân lớp, dễ thay thế repository mock bằng backend thật.",
    "Hoàn thiện luồng khách hàng từ onboarding/đăng nhập đến chọn sản phẩm, checkout và theo dõi đơn.",
    "Tích hợp lưu trữ cục bộ, bản đồ/vị trí và Firestore chat để thể hiện các năng lực ngoài UI cơ bản.",
    "Cung cấp bộ chức năng quản trị đủ cho demo CRUD, trạng thái đơn, doanh thu và tài khoản.",
]: add_number(doc, x, objective_num)

add_heading(doc, "3. Business Analysis / System Design", 1)
add_heading(doc, "3.1. Functional Requirements", 2)
add_table(doc, ["Nhóm tác nhân", "Yêu cầu chức năng", "Trạng thái thực tế"], [
    ["Khách/Người dùng", "Onboarding, đăng ký, đăng nhập, quên mật khẩu, ghi nhớ đăng nhập, sửa hồ sơ và địa chỉ", "Đã có UI/logic; runtime đang dùng MockAuthRepository"],
    ["Người mua", "Xem 28 sản phẩm, tìm kiếm, lọc theo danh mục, sắp xếp, xem chi tiết và sản phẩm liên quan", "Đã triển khai với JSON/repository mock"],
    ["Người mua", "Wishlist, giỏ hàng, tăng/giảm số lượng, kiểm tra tồn kho, phí vận chuyển và tổng tiền", "Wishlist: SQLite; cart: SharedPreferences"],
    ["Người mua", "Checkout, chọn địa chỉ, ghi chú, COD/Momo/VNPay, xác nhận, lịch sử và hủy đơn", "Đã triển khai; thanh toán online là luồng mô phỏng"],
    ["Người mua", "Thông báo, chat hỗ trợ, bản đồ cửa hàng và vị trí hiện tại", "Notification mock; chat Firestore/fallback mock; Google Maps + Geolocator"],
    ["Admin", "Dashboard, CRUD sản phẩm, trạng thái đơn, doanh thu, lọc/vô hiệu hóa người dùng", "Đã có UI/logic, dữ liệu mock/cục bộ"],
], [1500, 4770, 3090], font_size=8.5)

add_heading(doc, "3.2. Non-functional Requirements", 2)
add_bullet(doc, "Usability: Material 3, typography Outfit qua google_fonts, giao diện tiếng Việt, điều hướng tập trung và bố cục tối ưu cho màn hình dọc.")
add_bullet(doc, "Maintainability: tách Models - Repositories - ViewModels/Providers - Views; dependency injection repository tại App root.")
add_bullet(doc, "Reliability: repository chat tự chuyển sang mock khi Firebase không sẵn sàng; dữ liệu wishlist trên Web chuyển sang SharedPreferences vì SQLite không hỗ trợ.")
add_bullet(doc, "Security: Firestore rules yêu cầu xác thực, giới hạn chat theo userId và cho phép admin dựa trên trường role. Cần lưu ý app runtime hiện dùng MockAuthRepository nên mô hình xác thực production chưa khép kín.")
add_bullet(doc, "Performance: danh sách có phân trang 10 sản phẩm/trang, ảnh mạng dùng cached_network_image. Chưa có benchmark chứng minh SLA tải dưới 2 giây.")

add_heading(doc, "3.3. Application Architecture", 2)
add_p(doc, "Ứng dụng dùng Provider/ChangeNotifier kết hợp Repository Pattern. Dòng dữ liệu chính:")
add_table(doc, ["Presentation Layer", "State / Business Logic", "Data & External Services"], [
    ["Screens, widgets, Material routes", "Auth, Product, Cart, Checkout, Order, Notification, Chat, UserAdmin ViewModels; WishlistProvider", "Mock/SharedPreferences/SQLite repositories; Firebase Core & Firestore; Google Maps/Geolocator"],
], [2800, 3260, 3300], font_size=9)
add_bullet(doc, "App root khởi tạo MultiProvider với 9 provider/viewmodel và inject repository cụ thể.")
add_bullet(doc, "AppRoutes định nghĩa 28 named routes, bao phủ customer flow và admin flow.")
add_bullet(doc, "FirebaseService.initialize() quyết định dùng FirebaseChatRepository hay MockChatRepository.")

add_heading(doc, "3.4. Data Design", 2)
add_heading(doc, "3.4.1. SQLite - wishlist", 3)
add_table(doc, ["Cột", "Kiểu", "Ràng buộc / Ý nghĩa"], [
    ["id", "INTEGER", "PRIMARY KEY AUTOINCREMENT"],
    ["user_id", "INTEGER", "Định danh người dùng cục bộ; mặc định provider đang dùng user 1"],
    ["product_id", "TEXT", "UNIQUE; khóa đối chiếu sản phẩm"],
    ["product_name", "TEXT", "Tên hiển thị"],
    ["product_image", "TEXT", "URL ảnh chính"],
    ["price", "REAL", "Giá tại thời điểm lưu"],
    ["rating", "REAL", "Điểm đánh giá"],
    ["created_at", "TEXT", "Thời điểm thêm theo ISO-8601"],
], [1800, 1300, 6260], font_size=8.5)
add_heading(doc, "3.4.2. SharedPreferences", 3)
add_bullet(doc, "Cart: JSON list theo khóa cart_<userId>.")
add_bullet(doc, "Orders: danh sách OrderModel được serialize cục bộ.")
add_bullet(doc, "Authentication UX: remember_me và saved_email; onboarding_done cho trạng thái onboarding.")
add_bullet(doc, "Web wishlist fallback: kiencare_wishlist_user_<userId>.")
add_heading(doc, "3.4.3. Cloud Firestore - chat", 3)
add_bullet(doc, "Cấu trúc chats/{chatId}/messages/{messageId}, trong đó chatId = userId.")
add_bullet(doc, "Message yêu cầu senderId, senderName, text, isRead và timestamp; update chỉ cho phép trường isRead.")
add_bullet(doc, "Admin được xác định từ users/{uid}.role == admin trong security rules.")

add_heading(doc, "3.5. New Technologies", 2)
add_table(doc, ["Công nghệ", "Ứng dụng trong dự án"], [
    ["Firebase Core / Cloud Firestore", "Khởi tạo cloud service và realtime support chat có security rules"],
    ["Google Maps Flutter + Geolocator", "Hiển thị cửa hàng, marker, camera map và xác định vị trí người dùng"],
    ["sqflite", "Cơ sở dữ liệu wishlist cục bộ trên mobile"],
    ["SharedPreferences", "Lưu cart, orders, trạng thái onboarding và remember-me"],
    ["Provider", "Quản lý state và dependency injection ở cấp ứng dụng"],
    ["cached_network_image", "Cache và hiển thị ảnh sản phẩm từ mạng"],
], [2900, 6460], font_size=9)

add_heading(doc, "4. Development Requirements", 1)
add_heading(doc, "4.1. Implementation Details", 2)
add_p(doc, "Stack chính: Flutter SDK >= 3.3.0 < 4.0.0, Dart, Material 3, Provider 6.1.2, Firebase Core/Auth/Firestore, sqflite, SharedPreferences, Google Maps, Geolocator và các gói UI/utility trong pubspec.yaml.")
add_table(doc, ["Module", "Triển khai thực tế", "Ghi chú giới hạn"], [
    ["Authentication", "Có FirebaseAuthRepository đầy đủ và MockAuthRepository; app root hiện inject mock", "Tài khoản demo; chưa dùng Firebase Auth ở runtime mặc định"],
    ["Product", "28 laptop từ laptophub_products.json; fallback embedded JSON; search/filter/sort/paging", "Đường dẫn file JSON hard-code theo workspace; APK thường dùng embedded fallback"],
    ["Cart/Order", "ViewModel + SharedPreferences repository, tính subtotal, saving, ship fee và total", "Chưa đồng bộ cloud/đa thiết bị"],
    ["Payment", "COD/Momo/VNPay screens và trạng thái xử lý", "Không gọi cổng thanh toán thật"],
    ["Admin", "Dashboard, product/order/revenue/user screens", "Product/user/notification chủ yếu mock; thay đổi không phải backend production"],
    ["Chat", "Firestore realtime khi Firebase ready, mock fallback", "Cần tài khoản/role đồng nhất với Firestore để vận hành production"],
], [1700, 4940, 2720], font_size=8.2)

doc.add_page_break()
add_heading(doc, "4.2. Testing and Static Analysis", 2)
add_table(doc, ["Hạng mục", "Kết quả ngày 22/06/2026", "Đánh giá"], [
    ["flutter test", "11/11 tests passed", "PASS"],
    ["AuthViewModel unit tests", "8 ca: register, login/remember, initialize, logout, reset, error", "PASS"],
    ["Forgot password widget tests", "2 ca: email hợp lệ và email sai định dạng", "PASS"],
    ["App smoke test", "KienCareApp khởi tạo không crash", "PASS"],
    ["flutter analyze", "44 issues: 4 warning và 40 info/lint", "Cần cải thiện"],
], [2450, 4960, 1950], fills=[GREEN, GREEN, GREEN, GREEN, AMBER], font_size=8.8)
add_p(doc, "Các cảnh báo chính gồm unused imports, duplicate import và unused optional parameter; phần info chủ yếu là const, API deprecated và BuildContext qua async gap. Không có lỗi test, nhưng nên xử lý lint trước khi phát hành production.")

add_heading(doc, "4.3. Deployment", 2)
add_table(doc, ["Thông số", "Kết quả"], [
    ["Lệnh", "flutter build apk --release"],
    ["Trạng thái", "Build thành công ngày 22/06/2026"],
    ["Artifact", "build/app/outputs/flutter-apk/app-release.apk"],
    ["Kích thước", "58,899,840 bytes (Flutter hiển thị 56.2 MB)"],
    ["Ghi chú", "Build có cảnh báo Java source/target 8 và deprecated API từ dependency; không làm build thất bại"],
], [2100, 7260], font_size=9)

add_heading(doc, "5. Demo of Functions", 1)
demo_rows = [
    ("Splash & Onboarding", "Điều hướng lần đầu, lưu onboarding_done."),
    ("Authentication", "Đăng ký, đăng nhập, nhớ email, quên mật khẩu, logout; admin demo."),
    ("Home", "Banner, danh mục, promotion, sản phẩm nổi bật và hot deal."),
    ("Product List", "Grid, phân trang, tìm kiếm, lọc danh mục và sắp xếp."),
    ("Product Detail", "Thông số laptop, ảnh, rating, tồn kho, related products, add cart/wishlist."),
    ("Wishlist", "Thêm/xóa/xóa tất cả, optimistic UI; SQLite trên mobile."),
    ("Cart", "Tăng/giảm/xóa, giới hạn stock, tổng tiền, tiết kiệm, phí ship miễn phí từ 500.000đ."),
    ("Address & Checkout", "Chọn/thêm/sửa địa chỉ, ghi chú, COD/Momo/VNPay, đặt đơn."),
    ("Orders", "Lịch sử, chi tiết, trạng thái và hủy đơn phù hợp."),
    ("Notifications", "Danh sách, unread count, đọc từng/all; dữ liệu mock."),
    ("Chat", "Realtime Firestore khi sẵn sàng hoặc fallback mock."),
    ("Store Location", "Google Map, marker cửa hàng, quyền vị trí và vị trí hiện tại."),
    ("Profile", "Sửa hồ sơ, địa chỉ, lối vào admin cho tài khoản role admin."),
    ("Admin", "Dashboard, products CRUD, orders/status, revenue và users/status."),
]
add_table(doc, ["Màn hình/chức năng", "Nội dung demo"], demo_rows, [2300, 7060], font_size=8.5)

add_heading(doc, "5.1. Suggested Demo Script", 2)
demo_num = new_numbering(doc)
for x in [
    "Mở app, hoàn tất onboarding và đăng nhập tài khoản demo.",
    "Từ Home, lọc laptop Gaming, mở chi tiết, thêm wishlist và giỏ hàng.",
    "Trong Cart, thay đổi số lượng để kiểm tra tồn kho, subtotal và phí vận chuyển.",
    "Thêm/chọn địa chỉ, chọn phương thức thanh toán, đặt hàng và xem lịch sử đơn.",
    "Mở Notifications, Chat và Store Location để minh họa tích hợp dịch vụ.",
    "Đăng nhập admin, trình diễn dashboard, CRUD product, đổi trạng thái order và quản lý user.",
]: add_number(doc, x, demo_num)

add_heading(doc, "6. Conclusion and Discussion", 1)
add_heading(doc, "6.1. Strengths", 2)
add_bullet(doc, "Phạm vi chức năng rộng, bao phủ đầy đủ customer journey và admin journey cho một đồ án Flutter.")
add_bullet(doc, "Kiến trúc tách lớp và repository abstraction giúp thay nguồn dữ liệu mà ít ảnh hưởng UI.")
add_bullet(doc, "Có lưu trữ cục bộ thực tế, map/location và Firestore chat thay vì chỉ là màn hình tĩnh.")
add_bullet(doc, "11 automated tests pass và APK release build thành công.")
add_heading(doc, "6.2. Limitations", 2)
add_bullet(doc, "Auth runtime, product, notification và user-admin vẫn chủ yếu là mock; FirebaseAuthRepository chưa được inject mặc định.")
add_bullet(doc, "Thanh toán Momo/VNPay là mô phỏng, không có SDK/callback/chữ ký giao dịch thật.")
add_bullet(doc, "Cart/order dùng SharedPreferences, không phù hợp dữ liệu lớn, transaction phức tạp hoặc đồng bộ đa thiết bị.")
add_bullet(doc, "MockProductRepository dùng đường dẫn file tuyệt đối của máy phát triển trước khi fallback embedded JSON.")
add_bullet(doc, "44 vấn đề lint còn tồn tại; chưa có test cho cart calculation, order, wishlist SQLite, product filtering và Firestore integration.")
add_heading(doc, "6.3. Lessons Learned", 2)
add_bullet(doc, "Kết hợp ChangeNotifier/Provider với Repository Pattern và dependency injection.")
add_bullet(doc, "Thiết kế offline-first ở mức prototype bằng SQLite và SharedPreferences.")
add_bullet(doc, "Xử lý permission vị trí, Google Maps, Firebase initialization/fallback và Firestore security rules.")
add_heading(doc, "6.4. Future Improvements", 2)
future_num = new_numbering(doc)
for x in [
    "Chuyển app root sang FirebaseAuthRepository và lưu users/roles trong Firestore đồng nhất với security rules.",
    "Đưa products, orders, notifications và admin operations lên backend/Firestore; bỏ đường dẫn JSON tuyệt đối.",
    "Tích hợp payment gateway sandbox chính thức, xác minh callback phía server và trạng thái giao dịch.",
    "Bổ sung unit/widget/integration tests cho cart, checkout, wishlist, product và chat; xử lý toàn bộ lint.",
    "Bổ sung logging, crash reporting, analytics, pagination phía server và chiến lược cache/sync.",
]: add_number(doc, x, future_num)

add_heading(doc, "7. Contribution Table", 1)
add_p(doc, "Bảng sau phản ánh phân công theo vùng mã nguồn nổi bật trong lịch sử Git. Mỗi chủ đề vẫn có review chéo; tổng tỷ lệ theo từng hàng bằng 100%.")
add_table(doc, ["Topic", "Bùi Mạnh Thịnh", "Lê Trung Kiên", "Trần Đình Việt", "Trần Văn Phi"], [
    ["Case Study Analysis", "25%", "25%", "25%", "25%"],
    ["Business Analysis", "25%", "25%", "25%", "25%"],
    ["System Design", "30%", "30%", "25%", "15%"],
    ["Implementation", "25%", "25%", "25%", "25%"],
    ["Testing & Documentation", "25%", "25%", "25%", "25%"],
], [2360, 1750, 1750, 1750, 1750], font_size=8.5)

add_heading(doc, "8. References", 1)
refs = [
    "Flutter documentation - https://docs.flutter.dev/",
    "Provider package - https://pub.dev/packages/provider",
    "Firebase for Flutter - https://firebase.google.com/docs/flutter/setup",
    "Cloud Firestore - https://firebase.google.com/docs/firestore",
    "Firebase Security Rules - https://firebase.google.com/docs/rules",
    "sqflite - https://pub.dev/packages/sqflite",
    "SharedPreferences - https://pub.dev/packages/shared_preferences",
    "Google Maps Flutter - https://pub.dev/packages/google_maps_flutter",
    "Geolocator - https://pub.dev/packages/geolocator",
    "Nguồn nội bộ dự án: pubspec.yaml, lib/, test/, firestore.rules, laptophub_products.json và Git history (đối chiếu ngày 22/06/2026).",
]
for ref in refs: add_bullet(doc, ref)

doc.add_page_break()
add_heading(doc, "Phụ lục A. Ma trận đối chiếu với source code", 1)
add_table(doc, ["Khẳng định trong báo cáo", "Bằng chứng trong dự án"], [
    ["9 provider/viewmodel tại app root", "lib/app/app.dart - MultiProvider"],
    ["28 named routes", "lib/app/routes.dart"],
    ["28 sản phẩm JSON", "laptophub_products.json"],
    ["Wishlist SQLite", "lib/services/wishlist_database_service.dart"],
    ["Cart/Orders SharedPreferences", "lib/repositories/cart_repository.dart; order_repository.dart"],
    ["Auth mock runtime", "lib/app/app.dart inject MockAuthRepository"],
    ["Firestore chat + fallback", "lib/repositories/firebase_chat_repository.dart; lib/app/app.dart"],
    ["Chat security", "firestore.rules"],
    ["Google Maps/Geolocator", "lib/views/store_location/store_location_screen.dart"],
    ["11 tests pass", "test/*.dart và kết quả flutter test ngày 22/06/2026"],
    ["APK release 56.2 MB", "build/app/outputs/flutter-apk/app-release.apk"],
], [3850, 5510], font_size=8.5)

# Keep table rows together where practical and set core properties.
props = doc.core_properties
props.title = "KienCare / LaptopHub - Technical Project Report"
props.subject = "PRM393 Flutter Application Development"
props.author = "KienCare Project Team"
props.keywords = "Flutter, KienCare, LaptopHub, e-commerce, Firebase, SQLite, Provider"

doc.save(OUT)
print(OUT)
